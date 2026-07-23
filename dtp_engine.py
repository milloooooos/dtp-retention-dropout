# -*- coding: utf-8 -*-
"""
DTP 患者服务 · 自助分析引擎（留存率 / 脱落率A / 脱落率B / 跨表脱落原因）
=============================================================================
设计目标：把"销售底表 + 随访任务表"两份数据，统一算：
  1) 留存率 Cohort（按首购月分群 = 新患队列，追踪第 N 月仍在购药比例）
  2) 脱落率 A（滚动：基准月 M-2 有购药、观察窗 M-1∪M 无购药 → 脱落）
  3) 脱落率 B（累计沉默：末次购药距数据终点 > 3×用药间隔 → 真停药）
  4) 跨表关联：销售算出的脱落患者 → 随访表的脱落原因（细分类 + 一级分类：医生相关/患者相关/其他）
口径说明见 README。本模块被 run_dtp.py（本地批量）与 app.py（Streamlit）共用。
"""
import pandas as pd, numpy as np, re, os, json, io

# 版本号：用于确认 Streamlit Cloud 线上跑的是哪一份代码
APP_VERSION = '2026-07-23c'

# ---------------- 脱落原因清洗核心：复用 dtp-churn-analyzer 的纯逻辑模块 ----------------
# churn_logic.py 提供：CATEGORY_DISPLAY（分类代码→中文）/ LEVEL1_OF（一级：医生/患者/其他）
#   classify_reason()（四原则：精确优先→关键词兜底→医生相关→患者相关→其他→统一兜底10_其他）
#   status_to_group() / patient_level_detail() / extract_target_patients() 等。
# 本模块只在其上叠加「销售↔随访」跨表关联与双视角/覆盖率，不改 churn_logic 内部归因规则。
import churn_logic as cl

# ---------------- 品种与用药间隔（每盒天数，用于口径B阈值） ----------------
BRANDS = ['泰瑞沙', '优赫得', '利普卓', '英飞凡', '沃瑞沙', '凡舒卓', '荃科得']
# 每盒天数（= 一次用药间隔，口径B阈值 = 3 × 此值；也用于留存率覆盖口径）
DAYS_PER_BOX = {'泰瑞沙': 30, '利普卓': 14, '沃瑞沙': 7, '英飞凡': 28,
                '优赫得': 21, '凡舒卓': 28, '荃科得': 30}

# ================= 自丰富品牌映射表（通用名/别名 → 品牌） =================
# 说明：随访表写法为「品牌-(通用名)」、销售表写法为「通用名(品牌)」，两者都含品牌字，
# 但未来若某文件只给通用名（无品牌字），需靠映射表识别。本表会：
#   ① 从数据自动学习（凡是同时含品牌字的记录，抽出其通用名核心词→品牌，写回映射表）；
#   ② 对纯通用名记录用映射表回退匹配；③ 仍识别不出的写入 unresolved 供人工补录。
ALIAS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'brand_aliases.json')

def _seed_aliases():
    # 种子映射均来自本项目真实数据抽取，非主观假设
    return {
        '甲磺酸奥希替尼片': '泰瑞沙', '奥希替尼': '泰瑞沙',
        '注射用德曲妥珠单抗': '优赫得', '德曲妥珠单抗': '优赫得',
        '奥拉帕利片': '利普卓', '奥拉帕利': '利普卓', '奥拉帕尼': '利普卓',
        '度伐利尤单抗注射液': '英飞凡', '度伐利尤单抗': '英飞凡',
        '赛沃替尼片': '沃瑞沙', '赛沃替尼': '沃瑞沙',
        '本瑞利珠单抗注射液': '凡舒卓', '本瑞利珠单抗': '凡舒卓',
        '卡匹色替片': '荃科得', '卡匹色替': '荃科得',
    }

def load_alias_map():
    seed = _seed_aliases()
    if os.path.exists(ALIAS_FILE):
        try:
            with open(ALIAS_FILE, 'r', encoding='utf-8') as f:
                d = json.load(f)
            seed.update(d.get('aliases', {}))
            return {'aliases': seed, 'unresolved': list(d.get('unresolved', []))}
        except Exception:
            pass
    return {'aliases': seed, 'unresolved': []}

def save_alias_map(am):
    try:
        payload = {'aliases': am.get('aliases', {}),
                   'unresolved': sorted(set(am.get('unresolved', [])))}
        with open(ALIAS_FILE, 'w', encoding='utf-8') as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def _core_token(s):
    """去掉括号/连字符/空格等分隔符，得到核心药名串。"""
    return re.sub(r'[()（）\-—\s、,，/]+', '', str(s))

def learn_aliases(series, am):
    """从数据自动学习：凡记录同时含品牌字，抽出其通用名核心词 → 该品牌，写回映射表。"""
    aliases = am['aliases']
    for v in pd.Series(series).dropna().astype(str).unique():
        b = next((br for br in BRANDS if br in v), None)
        if not b:
            continue
        g = _core_token(v).replace(b, '')
        if g and g not in aliases:
            aliases[g] = b
    return am

def resolve_brand(x, am):
    """识别顺序：① 直接含品牌字 → ② 映射表通用名子串命中 → ③ 记入 unresolved 返回'其他'。"""
    s = str(x)
    b = next((br for br in BRANDS if br in s), None)
    if b:
        return b
    core = _core_token(s)
    for g, br in am['aliases'].items():
        if g and g in core:
            return br
    if s and s not in ('nan', 'NaN', '') and s not in am['unresolved']:
        am['unresolved'].append(s)
    return '其他'

# 模块级单例：整个进程共享一份映射表，随每次 load 学习并落盘
_ALIAS = load_alias_map()

# ---------------- 归因维度：仅用 churn_logic 的「细分类 + 一级分类」 ----------------
# 说明：本引擎只采用 churn_logic 给出的 10 个细类 + 一级分类（医生相关/患者相关/其他）
# 作为脱落原因的归因维度。**不引入「可控/不可控」这一层**（避免把医嘱停药等错误简化为
# 运营可干预，也避免用单一标签掩盖真实归因）。改进措施按「一级分类 + 细分原因」落责。

# 周期状态组(5组) → 三分类(脱落/风险/持续/其他)，供跨表分母与下游使用
_STATUS_TO_CLASS = {
    '停药——脱落': '脱落', '当期转渠道': '脱落', '随访失败': '脱落',
    '改变用药间隔时间': '风险', '按计划持续用药': '持续', '其他': '其他',
}
def status_class(g):
    if g is None or (isinstance(g, float) and pd.isna(g)):
        return '其他'
    return _STATUS_TO_CLASS.get(str(g).strip(), '其他')

# ---------------- 药房名归一化（销售表 ↔ 随访表 实体对齐） ----------------
# 两张表的药房命名常不一致（全称 vs 简称），用 pharm_map.json 人工维护别名对照；
# 复合患者键用「姓名||药房全称」，避免同人在不同药房被当成同一人（防重名串号）。
PHARM_MAP_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pharm_map.json')

def load_pharm_map():
    if os.path.exists(PHARM_MAP_FILE):
        try:
            with open(PHARM_MAP_FILE, 'r', encoding='utf-8') as f:
                d = json.load(f)
            return d.get('pharm_map', {}) or {}
        except Exception:
            pass
    return {}

def normalize_store(s, pharm_map=None):
    """药房名归一化：先查人工别名表，否则做「用于对齐」的归一（去空格/统一括号/
    大药房→药房/剔除国药控股·股份有限公司·有限公司·连锁 等法人壳词），便于销售表与随访表
    药房名对齐。不做激进模糊合并（不擅自把不同城市同名药房并为一处）。"""
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ''
    t = str(s).strip()
    if not t or t.lower() in ('nan', 'none', ''):
        return ''
    if pharm_map and t in pharm_map:
        return pharm_map[t]
    t = t.replace('（', '(').replace('）', ')')
    t = re.sub(r'\s+', '', t)
    t = t.replace('大药房', '药房')
    for tok in ('国药控股', '股份有限公司', '有限责任公司', '连锁有限公司', '连锁', '有限公司', '股份公司'):
        t = t.replace(tok, '')
    return t

def pharm_aligned(a, b):
    """两个归一化药房名是否对齐：互为子串即视为同一药房（用于置信度判定，非硬过滤）。"""
    if not a or not b:
        return False
    return a in b or b in a

def composite_key(name, store, pharm_map=None):
    """复合患者键：姓名 || 归一化药房全称（用非常见分隔符 || 防姓名含普通分隔符）。"""
    return f"{str(name).strip()}||{normalize_store(store, pharm_map)}"

# 脱落原因 → 行动建议（参考映射，按「一级分类 + 细分原因」组织，不使用可控/不可控标签）
ACTION_MAP = [
    ['患者相关', '经济负担、因经济原因永久停药、报销不方便',
     '推进慈善援助/医保报销落地、援助患者转介、支付方案宣导', '门店+AZ代表'],
    ['患者相关', '转竞争药房/其他药房/回医院/换城市/异地/配送不便',
     '院外联动锁客：首购即绑定复购提醒、就近取药/配送、竞品药房回流跟进', '门店运营'],
    ['患者相关', '联系不上/拒绝随访/未拨通/挂断/随访失败',
     '优化随访触达：多时段多次呼叫、企微/短信、家属协同', '门店随访'],
    ['患者相关', '观念不强/自主停药/不规律/自行延长间隔、对疾病与药物认知不足',
     '院内外联动：医生观念沟通 + 患者依从性教育 + 续方提醒', '门店+医生'],
    ['患者相关', '领取慈善援助药/赠药',
     '区分"真流失"与"援助分流"，记录援助渠道避免重复计脱落', '门店'],
    ['医生相关', '医嘱停药/遵医嘱改变用药间隔/医嘱换药',
     '与区域/处方医生沟通方案连续性，明确停药原因并评估是否可逆', '门店+医生+AZ代表'],
    ['医生相关', '因不良反应带来的剂量调整或停药',
     '联动医生做不良反应管理、剂量调整支持，降低非必要停药', '门店+医生'],
    ['其他', '其他原因（随访自由文本未归入上述细分）/去世等客观因素',
     '规范随访填写口径，提升原因可归类率；客观因素单列不干预', '门店+数据治理'],
]

# ---------------- 列名自动识别（兼容不同导出命名） ----------------
SALES_COL_ALIASES = {
    'brand_raw': ['商品名称', '产品名称', '药品名称'],
    'time': ['销售时间', '开单时间', '购药日期', '单据日期'],
    'qty': ['销售数量', '开单数量', '购药盒数', '数量'],
    'pid': ['患者ID', '患者编码', '会员号', '开票抬头'],
    'name': ['会员姓名', '患者姓名', '开票抬头', '患者'],
    'pharmacy': ['药房名称', '药店名称', '门店'],
    'hospital': ['医疗单位', '开方医院', '医院'],
    'doctor': ['处方医生', '医生'],
}
FOLLOWUP_COL_ALIASES = {
    'brand_raw': ['药品名称', '商品名称', '产品名称'],
    'name': ['患者', '患者姓名', '会员姓名'],
    'status': ['用药周期状态'],
    'create': ['创建日期'],
    'first_buy': ['患者首购日期', '首诊后首次用药时间'],
    'last_buy': ['末次购药日期'],
    'reason1': ['未按计划持续用药原因'],
    'reason2': ['本月未购药的原因'],
    'reason3': ['脱落/流失原因'],
    'reason_other': ['其他原因'],
    'task_type': ['任务类型'],
    'pharmacy': ['药店名称', '门店', '当前购药药店'],
}

def _pick(df, aliases):
    for a in aliases:
        if a in df.columns:
            return a
    return None

def brand_of(x):
    """兼容旧调用：走自丰富映射表识别（直接品牌字 → 通用名别名 → 其他）。"""
    return resolve_brand(x, _ALIAS)

# ---------------- 项目药房判定（TOP 清单） ----------------
TOP_COL_PATTERNS = ['TOP', '项目', '重点']

def load_pharmacy_tier(path):
    """读取项目药房 TOP 清单，返回 {药房全称: {品牌: '项目'/'非项目', '城市': str}}。
    支持列名：'对应的药店名称'/'药房名称'/'药店名称' + 各品种 TOP/项目/重点 列 + '城市'。"""
    df = pd.read_excel(path)
    df.columns = [str(c).strip() for c in df.columns]
    ph_col = next((c for c in df.columns if c in ['对应的药店名称', '药房名称', '药店名称', '门店']), None)
    city_col = next((c for c in df.columns if c in ['城市', '地市', '城市名']), None)
    if ph_col is None:
        return {}
    brand_cols = {}
    for c in df.columns:
        up = c.upper()
        for b in BRANDS:
            if b in c and any(p in up for p in TOP_COL_PATTERNS):
                brand_cols[b] = c
                break
    tier = {}
    for _, r in df.iterrows():
        ph = str(r[ph_col]).strip() if pd.notna(r[ph_col]) else None
        if not ph or ph in ('nan', 'NaN'):
            continue
        entry = {'城市': str(r[city_col]).strip() if city_col and pd.notna(r[city_col]) else ''}
        for b, col in brand_cols.items():
            v = r[col]
            if pd.notna(v) and str(v).strip() not in ('', 'nan', 'NaN'):
                txt = str(v).strip()
                entry[b] = '项目' if any(k in txt for k in ['项目', '重点', 'TOP', '是', 'Y', '√']) else '非项目'
        tier[ph] = entry
    return tier

def _normalize_pharmacy_name(s):
    """把常见简称统一成全称风格，便于和 TOP 清单匹配。"""
    if pd.isna(s):
        return ''
    return str(s).strip()

def apply_tier(sales, tier):
    """给销售底表加 角色/城市 列。如未提供 tier 且底表无角色列，默认全部为项目。"""
    sales = sales.copy()
    if '角色' not in sales.columns:
        sales['角色'] = '项目'
    if '城市' not in sales.columns:
        sales['城市'] = ''
    if not tier:
        return sales
    # 建立全称→tier 的查找（同时保留原名）
    norm_tier = {_normalize_pharmacy_name(k): v for k, v in tier.items()}
    roles = []
    cities = []
    for _, r in sales.iterrows():
        ph = _normalize_pharmacy_name(r.get('药房名称', ''))
        brand = r.get('品牌', '其他')
        entry = norm_tier.get(ph)
        # 尝试用子串匹配兜底
        if entry is None:
            for k, v in norm_tier.items():
                if k and (k in ph or ph in k):
                    entry = v
                    break
        if entry and brand in entry:
            roles.append(entry[brand])
        else:
            roles.append('非项目')
        cities.append(entry.get('城市', '') if entry else '')
    sales['角色'] = roles
    sales['城市'] = cities
    return sales

# ---------------- 加载销售底表 ----------------
def _resolve_sales_sheet(path, preferred):
    """销售表 sheet 容错：优先 preferred；缺失则按关键字(底表/销售/数据/明细/主)匹配，最后回退首个 sheet。"""
    if not isinstance(path, str):
        return preferred
    try:
        engine = 'xlrd' if path.lower().endswith('.xls') else None
        xf = pd.ExcelFile(path, engine=engine) if engine else pd.ExcelFile(path)
    except Exception:
        return preferred
    names = list(xf.sheet_names)
    if preferred in names:
        return preferred
    for kw in ('底表', '销售', '数据', '明细', '主'):
        for n in names:
            if kw in n:
                return n
    return names[0]


def _sales_sheet_names(path):
    try:
        engine = 'xlrd' if str(path).lower().endswith('.xls') else None
        xf = pd.ExcelFile(path, engine=engine) if engine else pd.ExcelFile(path)
        return ', '.join(xf.sheet_names)
    except Exception:
        return '（无法列出）'


def list_sales_sheets(path):
    """返回销售文件所有 sheet 名（list）；无法读取返回 []。供 app 下拉选择。"""
    if not isinstance(path, str):
        return []
    try:
        engine = 'xlrd' if path.lower().endswith('.xls') else None
        xf = pd.ExcelFile(path, engine=engine) if engine else pd.ExcelFile(path)
        return list(xf.sheet_names)
    except Exception:
        return []


def load_sales(path, sheet='底表', tier=None):
    if isinstance(path, pd.DataFrame):
        df = path
    else:
        real_sheet = _resolve_sales_sheet(path, sheet)
        try:
            if path.lower().endswith('.xls'):
                df = pd.read_excel(path, sheet_name=real_sheet, engine='xlrd')
            else:
                df = pd.read_excel(path, sheet_name=real_sheet)
        except Exception as e:
            raise ValueError(
                f"销售表读取失败（尝试 sheet={real_sheet!r}）：{e}。可用 sheet：{_sales_sheet_names(path)}"
            )
    df.columns = [str(c).strip() for c in df.columns]
    bcol = _pick(df, SALES_COL_ALIASES['brand_raw'])
    tcol = _pick(df, SALES_COL_ALIASES['time'])
    qcol = _pick(df, SALES_COL_ALIASES['qty'])
    pcol = _pick(df, SALES_COL_ALIASES['pid'])
    ncol = _pick(df, SALES_COL_ALIASES['name']) or pcol
    # 必需列校验（清晰报错，避免后续 KeyError）
    if bcol is None:
        raise ValueError(f"销售表缺少「商品名称/药品名称」类列。当前列名：{list(df.columns)}")
    if tcol is None:
        raise ValueError(f"销售表缺少「销售时间/购药日期」类列。当前列名：{list(df.columns)}")
    if pcol is None:
        raise ValueError(f"销售表缺少「患者ID/会员号」类列。当前列名：{list(df.columns)}")
    if qcol is None:
        raise ValueError(f"销售表缺少「销售数量/购药盒数」类列。当前列名：{list(df.columns)}")
    # 安全 rename（避免 同一列既作 pid 又作 name 时键冲突把 患者ID 改没）
    rename_map = {bcol: '品牌原始', tcol: '销售时间', qcol: '销售数量', pcol: '患者ID'}
    if ncol is not None and ncol != pcol:
        rename_map[ncol] = '患者姓名'
    df = df.rename(columns=rename_map).copy()
    if '患者姓名' not in df.columns:
        df['患者姓名'] = df['患者ID']
    df['销售时间'] = pd.to_datetime(df['销售时间'], errors='coerce')
    df = df.dropna(subset=['销售时间', '患者ID'])
    learn_aliases(df['品牌原始'], _ALIAS)                       # 自动学习
    df['品牌'] = df['品牌原始'].map(lambda x: resolve_brand(x, _ALIAS))
    save_alias_map(_ALIAS)                                      # 落盘丰富
    df['ym'] = df['销售时间'].dt.to_period('M')
    df['首购月'] = df.groupby('患者ID')['销售时间'].transform('min').dt.to_period('M')
    df['末次月'] = df.groupby('患者ID')['销售时间'].transform('max').dt.to_period('M')
    df['是否新患'] = df['ym'] == df['首购月']
    # 兼容 build_report 底表已有的药房/医院/医生列
    if '药房名称' not in df.columns:
        df['药房名称'] = df.get('门店', '未识别')
    if '医疗单位' not in df.columns:
        df['医疗单位'] = df.get('医院', '未识别')
    if '处方医生' not in df.columns:
        df['处方医生'] = df.get('医生', '')
    df = apply_tier(df, tier)
    return df

# ---------------- 加载随访任务表 ----------------
# 多原因列优先级：脱落/流失原因 > 未按计划持续用药原因 > 本月未购药的原因 > 脱落原因 > 其他原因
REASON_PRIORITY = ['脱落/流失原因', '未按计划持续用药原因', '本月未购药的原因', '脱落原因', '其他原因']

def _pick_reason_col(df):
    """按优先级返回第一个存在的脱落原因列；不存在返回 None。"""
    for c in REASON_PRIORITY:
        if c in df.columns:
            return c
    # 兜底：任意含「原因」「脱落」的列
    for c in df.columns:
        if '原因' in str(c) or '脱落' in str(c):
            return c
    return None

def _merge_reason(row, cols):
    """同一患者可能有多张随访任务，各填在不同原因列；取第一个非空作为该记录的原因原文。"""
    for c in cols:
        v = row.get(c)
        if pd.notna(v) and str(v).strip() not in ('', 'nan', 'NaN'):
            return str(v).strip()
    return None

def load_followup(paths, pharm_map=None):
    if isinstance(paths, (str,)):
        paths = [paths]
    frames = []
    for p in paths:
        try:
            frames.append(pd.read_excel(p, engine='xlrd'))
        except Exception:
            frames.append(pd.read_excel(p))
    df = pd.concat(frames, ignore_index=True, sort=False)
    if '任务编号' in df.columns:
        df = df.drop_duplicates(subset=['任务编号'], keep='first')
    df.columns = [str(c).strip() for c in df.columns]
    if pharm_map is None:
        pharm_map = load_pharm_map()
    bcol = _pick(df, FOLLOWUP_COL_ALIASES['brand_raw']) or '药品名称'
    ncol = _pick(df, FOLLOWUP_COL_ALIASES['name']) or '患者'
    scol = _pick(df, FOLLOWUP_COL_ALIASES['status']) or '用药周期状态'
    ccol = _pick(df, FOLLOWUP_COL_ALIASES['create']) or '创建日期'
    fcol = _pick(df, FOLLOWUP_COL_ALIASES['first_buy']) or '患者首购日期'
    lcol = _pick(df, FOLLOWUP_COL_ALIASES['last_buy']) or '末次购药日期'
    store_col = _pick(df, FOLLOWUP_COL_ALIASES['pharmacy'])
    df = df.rename(columns={bcol: '品牌原始', ncol: '患者姓名', scol: '用药周期状态',
                            ccol: '创建日期', fcol: '患者首购日期', lcol: '末次购药日期'}).copy()
    for c in ['创建日期', '患者首购日期', '末次购药日期']:
        if c in df.columns:
            df[c] = pd.to_datetime(df[c], errors='coerce')
    learn_aliases(df['品牌原始'], _ALIAS)                       # 自动学习
    df['品牌'] = df['品牌原始'].map(lambda x: resolve_brand(x, _ALIAS))
    save_alias_map(_ALIAS)                                      # 落盘丰富
    # ---- 脱落原因：多列合并 → churn_logic 细分类（医生优先四原则） ----
    rcol = _pick_reason_col(df)
    reason_cols = [c for c in REASON_PRIORITY if c in df.columns]
    if rcol:
        df['原因原始'] = df.apply(lambda r: _merge_reason(r, reason_cols), axis=1)
    else:
        df['原因原始'] = None
    df['脱落原因分类代码'] = df['原因原始'].map(lambda x: cl.classify_reason(x))
    df['脱落原因分类'] = df['脱落原因分类代码'].map(lambda c: cl.CATEGORY_DISPLAY.get(c) if c else None)
    df['一级分类'] = df['脱落原因分类代码'].map(lambda c: cl.LEVEL1_OF.get(c) if c else None)
    # ---- 用药周期状态 → 5 组 → 三分类 ----
    df['周期状态组'] = df['用药周期状态'].map(cl.status_to_group) if '用药周期状态' in df.columns else None
    df['状态类'] = df['周期状态组'].map(status_class)
    df['患者姓名'] = df['患者姓名'].astype(str).str.strip()
    # ---- 时间列（判定"最新一次"随访）→ _dt ----
    date_cols = [c for c in cl.DATE_COL_CANDIDATES if c in df.columns]
    if date_cols:
        df['_dt'] = df.apply(lambda r: cl.parse_row_datetime(r, date_cols), axis=1)
    else:
        df['_dt'] = pd.NaT
    # ---- 药房名归一化 + 复合患者键 ----
    if store_col:
        df = df.rename(columns={store_col: '药房原始'})
    else:
        df['药房原始'] = ''
    df['药房'] = df['药房原始'].map(lambda s: normalize_store(s, pharm_map))
    df['复合键'] = df.apply(lambda r: composite_key(r['患者姓名'], r['药房原始'], pharm_map), axis=1)
    return df

# ============================ 1. 留存率 Cohort ============================
def compute_retention(sales, max_k=12):
    """按首购月分群（=新患队列），追踪第 k 月仍在购药的比例。"""
    fp = sales.groupby('患者ID')['销售时间'].min()
    cohorts = fp.apply(lambda x: x.to_period('M'))
    sales = sales.copy()
    sales['cohort'] = sales['患者ID'].map(cohorts)
    max_month = sales['销售时间'].max().to_period('M')
    rows = []
    for cm in sorted(cohorts.unique()):
        grp = sales[sales['cohort'] == cm]
        base = set(grp['患者ID'].unique())
        n0 = len(base)
        rec = {'首购月': str(cm), '新患数': n0}
        for k in range(1, max_k + 1):
            target = cm + k
            if target > max_month:
                rec[f'M{k}留存%'] = None
                continue
            purch = set(sales[sales['ym'] == target]['患者ID'].unique())
            retained = base & purch
            rec[f'M{k}留存%'] = round(len(retained) / n0 * 100, 1) if n0 else None
        rows.append(rec)
    return pd.DataFrame(rows)

def compute_retention_by_brand(sales, max_k=12):
    out = []
    for b, g in sales.groupby('品牌'):
        if b == '其他':
            continue
        r = compute_retention(g, max_k)
        r.insert(0, '品牌', b)
        out.append(r)
    return pd.concat(out, ignore_index=True) if out else pd.DataFrame()

# -------- 留存率 口径2：结合说明书每盒天数的「覆盖期」判断 --------
def compute_retention_coverage(sales, max_k=12):
    """留存率(覆盖口径)：结合说明书『每盒天数×购买盒数』的药品覆盖期。
    每次购药覆盖 = 销售数量 × 每盒天数 天。患者在首购后第 k 个自然月，只要有
    任一购药的覆盖区间与该月有交集，即视为『仍在药品覆盖内=留存』。
    与口径1(仅看当月是否有购药)相比：一次囤多盒的患者在断购月仍算留存，
    曲线更贴近真实『在治』状态、更单调，剔除了多盒囤药造成的锯齿。"""
    s = sales.copy()
    s['dpb'] = s['品牌'].map(lambda b: DAYS_PER_BOX.get(b, 30))
    qty = pd.to_numeric(s['销售数量'], errors='coerce').fillna(1).clip(lower=1)
    s['cover_end'] = s['销售时间'] + pd.to_timedelta(qty * s['dpb'], unit='D')
    fp = s.groupby('患者ID')['销售时间'].min()
    cohorts = fp.apply(lambda x: x.to_period('M'))
    s['cohort'] = s['患者ID'].map(cohorts)
    max_month = s['销售时间'].max().to_period('M')
    rows = []
    for cm in sorted(cohorts.unique()):
        grp = s[s['cohort'] == cm]
        base = set(grp['患者ID'].unique())
        n0 = len(base)
        sub = s[s['患者ID'].isin(base)]
        rec = {'首购月': str(cm), '新患数': n0}
        for k in range(1, max_k + 1):
            target = cm + k
            if target > max_month:
                rec[f'M{k}留存%'] = None
                continue
            tstart = target.to_timestamp(how='start')
            tend = target.to_timestamp(how='end')
            covered = set(sub[(sub['销售时间'] <= tend) & (sub['cover_end'] >= tstart)]['患者ID'].unique())
            rec[f'M{k}留存%'] = round(len(covered) / n0 * 100, 1) if n0 else None
        rows.append(rec)
    return pd.DataFrame(rows)

def compute_retention_coverage_by_brand(sales, max_k=12):
    out = []
    for b, g in sales.groupby('品牌'):
        if b == '其他':
            continue
        r = compute_retention_coverage(g, max_k)
        r.insert(0, '品牌', b)
        out.append(r)
    return pd.concat(out, ignore_index=True) if out else pd.DataFrame()

def compute_retention_by_brand_pharmacy(sales, max_k=12):
    """留存率：分 品牌 × 药房。对每个(品牌,药房)子群跑 compute_retention（口径1：仅购药时间）。
    样本<3 的子群剔除，避免小样本噪声。用于定位『哪个药房哪个品种留存差』。"""
    pcol = '药房名称' if '药房名称' in sales.columns else None
    if not pcol:
        return pd.DataFrame()
    out = []
    for (b, ph), g in sales.groupby(['品牌', pcol]):
        if b == '其他' or g['患者ID'].nunique() < 3:
            continue
        r = compute_retention(g, max_k)
        r.insert(0, '药房', ph)
        r.insert(0, '品牌', b)
        out.append(r)
    return pd.concat(out, ignore_index=True) if out else pd.DataFrame()

# ============================ 2. 脱落率 A（滚动） ============================
def _monthly_dropout_A(sales):
    months = sorted(sales['ym'].unique())
    rows = []
    for i, m in enumerate(months):
        if i < 2:
            continue
        base_m, w1, w2 = months[i - 2], months[i - 1], months[i]
        base_p = set(sales[sales['ym'] == base_m]['患者ID'].unique())
        win_p = set(sales[sales['ym'].isin([w1, w2])]['患者ID'].unique())
        drop = base_p - win_p
        rows.append({'基准月': str(base_m), '基准患者数': len(base_p),
                     '脱落数': len(drop),
                     '脱落率A': round(len(drop) / len(base_p), 4) if base_p else None})
    return pd.DataFrame(rows)

def compute_dropout_A(sales):
    overall = _monthly_dropout_A(sales)
    res = {'整体_月度': overall}
    # 分品种
    brand_rows = []
    for b, g in sales.groupby('品牌'):
        if b == '其他':
            continue
        d = _monthly_dropout_A(g)
        d.insert(0, '品牌', b)
        brand_rows.append(d)
    if brand_rows:
        res['分品种_月度'] = pd.concat(brand_rows, ignore_index=True)
    # 分药房
    pharm_rows = []
    pcol = '药房名称' if '药房名称' in sales.columns else None
    if pcol:
        for ph, g in sales.groupby(pcol):
            d = _monthly_dropout_A(g)
            if d.empty:
                continue
            d.insert(0, '药房', ph)
            pharm_rows.append(d)
        if pharm_rows:
            res['分药房_月度'] = pd.concat(pharm_rows, ignore_index=True)
    return res

# ============================ 3. 脱落率 B（累计沉默） ============================
def compute_dropout_B(sales, mult=3, end=None):
    """患者级：末次购药距数据终点 > mult×用药间隔 → 真停药（脱落B）。
    同时标记 established：首购在「终点−mult×间隔」之前，才有足够观察期判断真停药；
    新近患者（首购太靠近终点）免除右删失，不应判定脱落。
    end：脱落判定终点；默认=None→用销售末次日期（全量累计口径），否则用所选时间窗终点（如 H1 2026-06-30）。"""
    if end is None:
        end = sales['销售时间'].max()
    last = sales.groupby('患者ID').agg(
        末次购药=('销售时间', 'max'), 品牌=('品牌', 'last'),
        首购=('销售时间', 'min'), 患者姓名=('患者姓名', 'last')).copy()
    if '药房名称' in sales.columns:
        last['药房'] = sales.groupby('患者ID')['药房名称'].last()
    last['距终点天'] = (end - last['末次购药']).dt.days
    last['阈值天'] = last['品牌'].map(lambda b: mult * DAYS_PER_BOX.get(b, 30))
    last['脱落B'] = last['距终点天'] > last['阈值天']
    last['观察窗起点'] = last['品牌'].map(lambda b: end - pd.Timedelta(days=mult * DAYS_PER_BOX.get(b, 30)))
    last['established'] = last['首购'] <= last['观察窗起点']
    return last.reset_index()

# ============================ 3b. 复购率 B（结合用药周期 / 首购→二购 on-cycle） ============================
def compute_repurchase_B(sales, mult=3, end=None):
    """患者级复购率B（结合用药周期）。
    定义：每个首购患者，若其第二次购药距首购 ≤ mult×用药间隔(首购品种) → 按周期复购(on_cycle_B)。
    右删失(established)：首购在「终点−mult×间隔」之后，观察期不足，尚无法判定二购是否on-cycle
        → 计入「已观察」分母时剔除（与脱落率B的 established 逻辑一致）。
    end：判定终点，默认销售末次日期；run_analysis 传窗口终点(wend)以与脱落率B口径一致。
    返回每患者一行：首购/首购品牌/二购/二购gap天/阈值天/on_cycle_B/established。"""
    if end is None:
        end = sales['销售时间'].max()
    df = sales.sort_values(['患者ID', '销售时间']).groupby('患者ID').agg(
        首购=('销售时间', 'first'),
        首购品牌=('品牌', 'first'),
        二购=('销售时间', lambda s: s.iloc[1] if len(s) > 1 else pd.NaT),
    ).reset_index()
    df['二购gap天'] = (df['二购'] - df['首购']).dt.days
    df['阈值天'] = df['首购品牌'].map(lambda b: mult * DAYS_PER_BOX.get(b, 30))
    df['on_cycle_B'] = df['二购'].notna() & (df['二购gap天'] <= df['阈值天'])
    df['观察窗起点'] = df['首购品牌'].map(lambda b: end - pd.Timedelta(days=mult * DAYS_PER_BOX.get(b, 30)))
    df['established'] = df['首购'] <= df['观察窗起点']
    return df


def repurchase_B_by_brand(sales, mult=3, end=None):
    """复购率B 分品种：全量(首购患者) + 已观察(剔除右删失) 两个口径。
    镜像 compute_dropout_B → dropout_B_by_brand 的结构。"""
    df = compute_repurchase_B(sales, mult, end)
    if df.empty:
        return pd.DataFrame()
    rows = []
    for b, g in df.groupby('首购品牌'):
        total = len(g)
        on_cycle = int(g['on_cycle_B'].sum())
        est = g[g['established']]
        est_total = len(est)
        est_on = int(est['on_cycle_B'].sum())
        rows.append({
            '品种': b,
            '复购率B%': round(on_cycle / total * 100, 1) if total else '',
            '按周期复购数': on_cycle,
            '首购患者数': total,
            '复购率B_已观察%': round(est_on / est_total * 100, 1) if est_total else '',
            '已观察_按周期复购数': est_on,
            '已观察患者数': est_total,
        })
    out = pd.DataFrame(rows)
    if out.empty:
        return out
    return out.sort_values('复购率B_已观察%', ascending=False)


# ============================ 4. 跨表关联：销售脱落 ↔ 随访原因 ============================
def crossref_brand(sales_last, followup, wstart=None, wend=None):
    """品牌层面关联：销售各品种脱落B人数 ↔ 随访各品种脱落原因构成（按一级分类：医生相关/患者相关/其他）。
    wstart/wend：仅统计 _dt 落在窗内的随访记录。"""
    s = sales_last.groupby('品牌')['脱落B'].agg(脱落患者数='sum', 患者总数='count')
    s['销售脱落率B%'] = (s['脱落患者数'] / s['患者总数'] * 100).round(1)
    fd = _reason_windowed(followup, wstart, wend)
    fd = fd[fd['状态类'] == '脱落'].copy()
    lvl_cols = ['一、医生相关原因', '二、患者相关原因', '三、其他原因']
    if fd.empty:
        reason = pd.DataFrame(columns=['品牌'] + [c + '数' for c in lvl_cols])
    else:
        grp = fd.groupby('品牌')['一级分类'].value_counts().unstack(fill_value=0)
        for c in lvl_cols:
            if c not in grp.columns:
                grp[c] = 0
        reason = grp.reset_index()
        reason = reason.rename(columns={c: c + '数' for c in lvl_cols})
    out = s.reset_index().merge(reason, on='品牌', how='left')
    return out

def _reason_windowed(followup, wstart=None, wend=None):
    """按时间窗过滤随访记录：仅保留 _dt 落在 [wstart, wend] 的记录（None 表示不过滤）。"""
    if wstart is not None and wend is not None and '_dt' in followup.columns:
        return followup[(followup['_dt'] >= wstart) & (followup['_dt'] <= wend)].copy()
    return followup.copy()

def dropout_reason_distribution(followup, wstart=None, wend=None):
    """脱落原因细分类分布：基于『有记载脱落原因』的随访记录（原因原始非空→已归类）。
    可选 wstart/wend：仅统计 _dt 落在所选时间窗内的随访记录（如 H1 2026）。
    返回 (dist[细类], lvl1[一级汇总], meta[覆盖])。归因维度仅用「细分类 + 一级分类」，不含可控/不可控。"""
    fw = _reason_windowed(followup, wstart, wend)
    sub = fw[fw['脱落原因分类代码'].notna()].copy()
    total = len(sub)
    rows = []
    for code, name in cl.CATEGORY_DISPLAY.items():
        cnt = int((sub['脱落原因分类代码'] == code).sum())
        if cnt == 0:
            continue
        rows.append({'一级分类': cl.LEVEL1_OF[code], '脱落原因细类': name,
                     '记录数': cnt, '占窗内%': round(cnt / total * 100, 1) if total else 0})
    dist = pd.DataFrame(rows).sort_values('记录数', ascending=False).reset_index(drop=True)
    lvl1 = dist.groupby('一级分类')['记录数'].sum().reset_index().sort_values('记录数', ascending=False)
    lvl1['占窗内%'] = (lvl1['记录数'] / total * 100).round(1) if total else 0
    total_records = len(fw)
    no_reason = int((fw['脱落原因分类代码'].isna()).sum())
    meta = {'随访窗内总记录': total_records, '窗内有原因记录': total, '窗内无原因记录': no_reason,
            '窗内无原因占比%': round(no_reason / total_records * 100, 1) if total_records else 0}
    return dist, lvl1, meta

def dropout_reason_detail_by_brand(followup, wstart=None, wend=None):
    """脱落原因细分类分布『按品种拆分汇总』：每个品种单独统计原因构成（窗内）。
    返回长表：品牌, 一级分类, 脱落原因细类, 记录数, 占该品种%。便于 Excel 按品种透视。"""
    fw = _reason_windowed(followup, wstart, wend)
    sub = fw[fw['脱落原因分类代码'].notna()].copy()
    if sub.empty:
        return pd.DataFrame(columns=['品牌', '一级分类', '脱落原因细类', '记录数', '占该品种%'])
    rows = []
    for b, g in sub.groupby('品牌'):
        n = len(g)
        for code, name in cl.CATEGORY_DISPLAY.items():
            cnt = int((g['脱落原因分类代码'] == code).sum())
            if cnt == 0:
                continue
            rows.append({'品牌': b, '一级分类': cl.LEVEL1_OF[code], '脱落原因细类': name,
                         '记录数': cnt, '占该品种%': round(cnt / n * 100, 1)})
    out = pd.DataFrame(rows).sort_values(['品牌', '记录数'], ascending=[True, False]).reset_index(drop=True)
    return out

def crossref_patient(sales_last, followup, pharm_map=None, wstart=None, wend=None):
    """销售脱落患者 ↔ 随访原因。两表无共同患者主键，且药房命名体系不同
    （销售:「成都晟德药房」/ 随访:「四川省晟德药房有限公司」），故：
      · 主匹配键 = 患者姓名（同一人对齐的基础）；
      · 药房作为『软校验/置信度』：若随访「最近一次有原因」记录的药房与销售药房归一化后对齐
        → 标记『姓名+药房(对齐)』(高置信)；否则 → 『姓名匹配(药房未对齐,近似)』(中低置信)。
    对每位随访患者取『双视角』：最新一次记录(当前状态) 与 最近一次有原因记录(有记载的脱落原因)。
    覆盖率：分母=销售脱落B患者；分子=随访表按姓名匹配上；区分『有记载原因』与『有记录但原因未填』。
    低匹配率(<30%)→attrs['low_match']=True，提示可能传错底表(药品/项目/时间段不同)。"""
    if pharm_map is None:
        pharm_map = load_pharm_map()
    ch = sales_last[sales_last['脱落B']].copy()
    if ch.empty:
        out = pd.DataFrame()
        out.attrs.update({'matched': 0, 'total': 0, 'with_reason': 0,
                          'no_reason_record': 0, 'unmatched': 0, 'match_rate': 0.0, 'low_match': False})
        return out
    ch['销售药房'] = ch.apply(lambda r: r.get('药房') or r.get('药房名称') or '', axis=1)
    ch['销售药房_norm'] = ch['销售药房'].map(lambda s: normalize_store(s, pharm_map))

    fcols = ['患者姓名', '药房', '品牌', '原因原始', '脱落原因分类代码',
             '脱落原因分类', '一级分类', '周期状态组', '_dt']
    fsub = followup[[c for c in fcols if c in followup.columns]].copy()
    fsub = _reason_windowed(fsub, wstart, wend)
    fsub['药房_norm'] = fsub['药房'].map(lambda s: normalize_store(s, pharm_map))
    fu_groups = {}
    for name, g in fsub.groupby('患者姓名'):
        gg = g.sort_values('_dt', ascending=False)
        last = gg.iloc[0]
        last_reason = last.get('原因原始')
        last_reason = '' if (last_reason is None or pd.isna(last_reason)) else str(last_reason).strip()
        last_code = cl.classify_reason(last_reason) if last_reason else None
        gr = gg[gg['脱落原因分类代码'].notna()]
        if len(gr) > 0:
            lr = gr.iloc[0]
            lr_reason = '' if pd.isna(lr.get('原因原始')) else str(lr.get('原因原始')).strip()
            lr_code = lr['脱落原因分类代码']
            lr_name = lr.get('脱落原因分类')
            lr_lvl = lr.get('一级分类')
            lr_pharm_norm = normalize_store(lr.get('药房', ''), pharm_map)
            has_reason = True
        else:
            lr_reason = ''; lr_code = None; lr_name = None; lr_lvl = None
            lr_pharm_norm = ''; has_reason = False
        fu_groups[name] = {
            'fu_最新周期状态': last.get('周期状态组'),
            'fu_最新一次原文': last_reason,
            'fu_最新一次分类': cl.CATEGORY_DISPLAY.get(last_code) if last_code else None,
            'fu_最近有原因原文': lr_reason,
            'fu_最近有原因分类': lr_name,
            'fu_最近有原因一级': lr_lvl,
            'fu_最近有原因药房_norm': lr_pharm_norm,
            'fu_药房集合_norm': set(fsub[fsub['患者姓名'] == name]['药房_norm'].dropna().unique()),
            'fu_有原因记录': has_reason,
        }
    rows = []
    matched = with_reason = no_reason_record = 0
    for _, r in ch.iterrows():
        name = r['患者姓名']
        sph = r['销售药房_norm']
        m = fu_groups.get(name)
        rec = {'患者姓名': name, '品牌': r['品牌'], '药房': r['销售药房'],
               '首购': r['首购'], '末次购药': r['末次购药'], '距终点天': int(r['距终点天']),
               'established': bool(r['established']), '关联方式': '姓名匹配(近似)'}
        if m:
            matched += 1
            rec['fu_最新周期状态'] = m['fu_最新周期状态']
            rec['最新一次原因(随访)'] = m['fu_最新一次原文']
            rec['最新一次分类'] = m['fu_最新一次分类']
            rec['脱落原因(随访)'] = m['fu_最近有原因原文']
            rec['脱落原因分类'] = m['fu_最近有原因分类']
            rec['一级分类'] = m['fu_最近有原因一级']
            rec['fu_有原因记录'] = m['fu_有原因记录']
            if m['fu_有原因记录']:
                with_reason += 1
                aligned = bool(sph) and (pharm_aligned(sph, m['fu_最近有原因药房_norm'])
                                         or any(pharm_aligned(sph, p) for p in m['fu_药房集合_norm']))
                rec['关联方式'] = '姓名+药房(对齐)' if aligned else '姓名匹配(药房未对齐,近似)'
            else:
                no_reason_record += 1
                rec['关联方式'] = '有随访记录但原因未填'
        else:
            rec['fu_最新周期状态'] = None
            rec['最新一次原因(随访)'] = None
            rec['最新一次分类'] = None
            rec['脱落原因(随访)'] = None
            rec['脱落原因分类'] = None
            rec['一级分类'] = None
            rec['fu_有原因记录'] = False
            rec['关联方式'] = '无随访匹配'
        rows.append(rec)
    out = pd.DataFrame(rows)
    total = len(ch)
    unmatched = total - matched
    match_rate = matched / total if total else 0.0
    out.attrs['matched'] = matched
    out.attrs['total'] = total
    out.attrs['with_reason'] = with_reason
    out.attrs['no_reason_record'] = no_reason_record
    out.attrs['unmatched'] = unmatched
    out.attrs['match_rate'] = match_rate
    out.attrs['low_match'] = (total >= 5 and match_rate < 0.3)
    return out

def suggest_pharm_map(sales, followup, pharm_map=None):
    """自动发现销售药房 ↔ 随访门店 的对齐候选，写回 pharm_map.auto.json（不覆盖用户维护的 pharm_map.json）。
    基于归一化后互为子串对齐；供人工确认后抄入 pharm_map.json。"""
    if pharm_map is None:
        pharm_map = load_pharm_map()
    try:
        sals = set(normalize_store(s, pharm_map) for s in sales['药房名称'].dropna().unique())
        fus = set(normalize_store(s, pharm_map) for s in followup['门店'].dropna().unique()) if '门店' in followup.columns \
            else set(normalize_store(s, pharm_map) for s in followup['药房原始'].dropna().unique())
        sugg = {}
        for sp in sals:
            for fp in fus:
                if sp and fp and (sp in fp or fp in sp) and sp != fp:
                    sugg[sp] = fp
                    break
        auto_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pharm_map.auto.json')
        with open(auto_file, 'w', encoding='utf-8') as f:
            json.dump({'pharm_map_suggestion': sugg}, f, ensure_ascii=False, indent=2)
        return sugg
    except Exception:
        return {}

def crossref_coverage_summary(cp):
    """覆盖率概览表（用于展示与低匹配报警）。"""
    total = cp.attrs.get('total', 0)
    matched = cp.attrs.get('matched', 0)
    with_reason = cp.attrs.get('with_reason', 0)
    no_reason_record = cp.attrs.get('no_reason_record', 0)
    unmatched = cp.attrs.get('unmatched', 0)
    match_rate = cp.attrs.get('match_rate', 0.0)
    rows = [
        {'指标': '销售脱落B患者(分母)', '人数': total},
        {'指标': '随访表按(姓名+药房)匹配上', '人数': matched,
         '占分母%': round(matched / total * 100, 1) if total else 0},
        {'指标': '  其中·有记载脱落原因', '人数': with_reason,
         '占分母%': round(with_reason / total * 100, 1) if total else 0},
        {'指标': '  其中·有随访记录但原因未填(空原因≠无脱落)', '人数': no_reason_record,
         '占分母%': round(no_reason_record / total * 100, 1) if total else 0},
        {'指标': '随访表无匹配(未随访/重名/传错底表)', '人数': unmatched,
         '占分母%': round(unmatched / total * 100, 1) if total else 0},
    ]
    df = pd.DataFrame(rows)
    df.attrs['match_rate'] = match_rate
    df.attrs['low_match'] = cp.attrs.get('low_match', False)
    return df

# ============================ 5. DOT / 新患 / 复购 / 维度分析 ============================
def resolve_window(sales, preset='H1_2026', custom_start=None, custom_end=None):
    """统一时间窗解析。返回 (wstart, wend, wstart_prev, wend_prev)。
    preset:
      'H1_2026'  : 本期 2026-01-01~2026-06-30，同比 2025-01-01~2025-06-30（半年度复盘口径）
      'roll1y'   : 回滚1年（本期=末次数据日往前1年，同比再往前1年），与DOT历史口径一致
      'full'     : 全量累计（起点=首笔销售，终点=末次销售），无同比
      'custom'   : 自定义区间，custom_start/custom_end 任意日期（支持跨年/不规则长度，
                   如 2026-01-01~2026-07-31、或 2027 全年等）；同比=同区间回退1年
    注意：'H1_2026' 的 wend 即脱落率B的判定终点；原因分布只统计 _dt 落在该窗内的随访记录。
    """
    smax = sales['销售时间'].max()
    smin = sales['销售时间'].min()
    if preset == 'custom':
        try:
            wstart = pd.Timestamp(custom_start)
            wend = pd.Timestamp(custom_end)
        except Exception:
            wstart, wend = pd.Timestamp(2026, 1, 1), pd.Timestamp(2026, 6, 30)
        if wend < wstart:  # 防止起止颠倒
            wstart, wend = wend, wstart
        wstart_prev = pd.Timestamp(wstart.year - 1, wstart.month, wstart.day)
        wend_prev = pd.Timestamp(wend.year - 1, wend.month, wend.day)
        return wstart, wend, wstart_prev, wend_prev
    if preset == 'roll1y':
        end = pd.Timestamp(smax.year, smax.month, 1) + pd.offsets.MonthEnd(0)
        start = (end - pd.DateOffset(years=1)) + pd.Timedelta(days=1)
        start = pd.Timestamp(start.year, start.month, 1)
        end_prev = start - pd.Timedelta(days=1)
        start_prev = (end_prev - pd.DateOffset(years=1)) + pd.Timedelta(days=1)
        start_prev = pd.Timestamp(start_prev.year, start_prev.month, 1)
        return start, end, start_prev, end_prev
    if preset == 'full':
        end = pd.Timestamp(smax.year, smax.month, 1) + pd.offsets.MonthEnd(0)
        start = pd.Timestamp(smin.year, smin.month, 1)
        return start, end, None, None
    # 默认 H1_2026（同比双窗共用同一组日期，仅报告展示差异）
    wstart = pd.Timestamp(2026, 1, 1)
    wend = pd.Timestamp(2026, 6, 30)
    wstart_prev = pd.Timestamp(2025, 1, 1)
    wend_prev = pd.Timestamp(2025, 6, 30)
    return wstart, wend, wstart_prev, wend_prev

def _dot_snap(sub, wstart, wend, pids=None):
    """DOT = 窗口内购买总盒数 / 去重患者数（盒/人，不换算月数）。"""
    if pids is not None:
        sub = sub[sub['患者ID'].isin(pids)]
    win = sub[(sub['销售时间'] >= wstart) & (sub['销售时间'] <= wend)]
    if win.empty:
        return np.nan
    return win.groupby('患者ID')['销售数量'].sum().mean()


def _dot_window(wend):
    """DOT 固定测量窗：窗口终点前推1年（盒/人，不按月归一）。
    无论分析窗长短（H1/自定义7个月/全年），DOT 一律在「[wend-1年+1天, wend]」
    上测算，保证跨分析窗可比，也对齐临床 DOT 习惯与本项目的「标准 DOT 口径」。
    返回 (dot_start, dot_end)；wend 为 None 时返回 (None, None)。
    """
    if wend is None:
        return None, None
    dot_start = (wend - pd.DateOffset(years=1)) + pd.Timedelta(days=1)
    return dot_start, wend

def _rolling_repurchase(sub, months, fm):
    """滚动复购率：M-2 有购药的患者中，M-1∪M 仍有购药的占比，按月平均。"""
    pm = sub.groupby(['患者ID', 'ym']).size().reset_index()
    have = set(sub['ym'].unique())
    rates = []
    for m in months:
        if (m - 2) not in have or m not in have:
            continue
        bp = set(pm[pm['ym'] == m - 2]['患者ID'])
        if not bp:
            continue
        ret = set(pm[pm['ym'].isin([m - 1, m])]['患者ID'])
        rates.append(len(bp & ret) / len(bp))
    return np.mean(rates) if rates else np.nan

def dot_decomposition(sales, wstart, wend, wstart_prev=None, wend_prev=None):
    """DOT 分解：全量 / 老患 / 新患窗口 / 重叠患者，同比。窗口由 resolve_window 给定。"""
    if wstart is None:
        wstart = sales['销售时间'].min()
    if wend is None:
        wend = sales['销售时间'].max()
    fm = sales.groupby('患者ID')['销售时间'].min()
    old = set(fm[fm < wstart].index)
    new = set(fm[(fm >= wstart) & (fm <= wend)].index)
    active = sales[(sales['销售时间'] >= wstart) & (sales['销售时间'] <= wend)]['患者ID'].unique()
    new = new & set(active)
    has_prev = wstart_prev is not None and wend_prev is not None
    if has_prev:
        old_prev = set(fm[fm < wstart_prev].index)
        new_prev = set(fm[(fm >= wstart_prev) & (fm <= wend_prev)].index)
        active_prev = sales[(sales['销售时间'] >= wstart_prev) & (sales['销售时间'] <= wend_prev)]['患者ID'].unique()
        new_prev = new_prev & set(active_prev)
        overlap = set(active) & set(active_prev)
        months_prev = list(pd.period_range(wstart_prev, wend_prev, freq='M'))
    else:
        old_prev = new_prev = overlap = set()
        months_prev = []
    months = list(pd.period_range(wstart, wend, freq='M'))
    # DOT 固定测量窗：窗口终点前推1年（与医生/医院/药房 DOT 口径一致，跨窗可比）
    dot_s, dot_e = _dot_window(wend)
    dot_s_prev, dot_e_prev = _dot_window(wend_prev) if has_prev else (None, None)
    rows = []
    for b in BRANDS:
        sub = sales[sales['品牌'] == b]
        if sub.empty:
            continue
        d_all = _dot_snap(sub, dot_s, dot_e)
        d_all_prev = _dot_snap(sub, dot_s_prev, dot_e_prev) if has_prev else np.nan
        d_old = _dot_snap(sub, dot_s, dot_e, old)
        d_old_prev = _dot_snap(sub, dot_s_prev, dot_e_prev, old_prev) if has_prev else np.nan
        d_new = _dot_snap(sub, dot_s, dot_e, new)
        d_new_prev = _dot_snap(sub, dot_s_prev, dot_e_prev, new_prev) if has_prev else np.nan
        d_ov = _dot_snap(sub, dot_s, dot_e, overlap)
        d_ov_prev = _dot_snap(sub, dot_s_prev, dot_e_prev, overlap) if has_prev else np.nan
        r_all = _rolling_repurchase(sub, months, fm)
        r_all_prev = _rolling_repurchase(sub, months_prev, fm) if months_prev else np.nan
        r_old = _rolling_repurchase(sub[sub['患者ID'].isin(old)], months, fm)
        r_old_prev = _rolling_repurchase(sub[sub['患者ID'].isin(old_prev)], months_prev, fm) if has_prev else np.nan
        rows.append({
            '品种': b, '患者数': sub['患者ID'].nunique(),
            'DOT_本期_全': round(d_all, 2) if pd.notna(d_all) else '',
            'DOT_同比_全': round(d_all_prev, 2) if pd.notna(d_all_prev) else '',
            'DOT全_Δ': round(d_all - d_all_prev, 2) if pd.notna(d_all) and pd.notna(d_all_prev) else '',
            'DOT_本期_老患': round(d_old, 2) if pd.notna(d_old) else '',
            'DOT_同比_老患': round(d_old_prev, 2) if pd.notna(d_old_prev) else '',
            'DOT老患_Δ': round(d_old - d_old_prev, 2) if pd.notna(d_old) and pd.notna(d_old_prev) else '',
            'DOT_本期_新患窗口': round(d_new, 2) if pd.notna(d_new) else '',
            'DOT_同比_新患窗口': round(d_new_prev, 2) if pd.notna(d_new_prev) else '',
            'DOT_本期_重叠患者': round(d_ov, 2) if pd.notna(d_ov) else '',
            'DOT_同比_重叠患者': round(d_ov_prev, 2) if pd.notna(d_ov_prev) else '',
            '复购率_本期_全': round(r_all, 3) if pd.notna(r_all) else '',
            '复购率_同比_全': round(r_all_prev, 3) if pd.notna(r_all_prev) else '',
            '复购率_本期_老患': round(r_old, 3) if pd.notna(r_old) else '',
            '复购率_同比_老患': round(r_old_prev, 3) if pd.notna(r_old_prev) else '',
        })
    return pd.DataFrame(rows)

def new_patient_monthly(sales):
    """新患月度趋势（按首购品种、首购月）。"""
    fm = sales.groupby('患者ID').agg(首购月=('ym', 'min'), 首购品种=('品牌', lambda x: x.iloc[0]))
    out = pd.DataFrame()
    for b in BRANDS:
        s = fm[fm['首购品种'] == b]
        vc = s['首购月'].value_counts().sort_index()
        out[b] = vc
    out.index.name = '月份'
    out = out.fillna(0).astype(int)
    # 取最近 18 个月
    min_m = out.index.min()
    if pd.notna(min_m):
        out = out[out.index >= (out.index.max() - 17)]
    return out.reset_index()

def new_patient_pharmacy_decline(sales, scope, wstart, wend, wstart_prev=None, wend_prev=None):
    """项目药房新患同比变化（本期 vs 上期，按品种+药房）。"""
    if wstart is None:
        wstart = sales['销售时间'].min()
    if wend is None:
        wend = sales['销售时间'].max()
    has_prev = wstart_prev is not None and wend_prev is not None
    fm = sales.groupby('患者ID').agg(首购月=('ym', 'min'), 首购品种=('品牌', lambda x: x.iloc[0]),
                                      首购药房=('药房名称', lambda x: x.iloc[0]))
    proj = sales[sales['角色'] == '项目'] if (scope == '项目' and '角色' in sales.columns) else sales
    rows = []
    for b in BRANDS:
        s = fm[fm['首购品种'] == b]
        for ph, g in s.groupby('首购药房'):
            if scope == '项目' and '角色' in sales.columns and role_of(ph, b, sales) != '项目':
                continue
            n_now = int(((g['首购月'] >= wstart.to_period('M')) & (g['首购月'] <= wend.to_period('M'))).sum())
            n_prev = int(((g['首购月'] >= wstart_prev.to_period('M')) & (g['首购月'] <= wend_prev.to_period('M'))).sum()) if has_prev else 0
            if (not has_prev) and n_now == 0:
                continue
            if has_prev and n_prev == 0 and n_now == 0:
                continue
            rows.append({'品种': b, '药房': ph, '城市': sales[sales['药房名称'] == ph]['城市'].dropna().iloc[0] if '城市' in sales.columns and not sales[sales['药房名称'] == ph]['城市'].dropna().empty else '',
                         '新患_本期': n_now, '新患_上期': n_prev, '变化': n_now - n_prev,
                         '同比': round((n_now - n_prev) / n_prev, 2) if n_prev else np.nan})
    return pd.DataFrame(rows).sort_values(['品种', '同比'])

def role_of(ph, b, sales):
    """从 sales 的角色列反查药房在某品种下的角色（项目/非项目）。"""
    sub = sales[(sales['药房名称'] == ph) & (sales['品牌'] == b)]
    if sub.empty or '角色' not in sub.columns:
        return '项目'
    return sub['角色'].iloc[0]

def old_patient_multi_box(sales, scope, wstart=None, wend=None, wstart_prev=None, wend_prev=None):
    """老患购药行为：单次购买盒数分布、购药间隔。"""
    if wstart is None or wend is None:
        wstart, wend, wstart_prev, wend_prev = resolve_window(sales, 'H1_2026')
    fm = sales.groupby('患者ID')['销售时间'].min()
    old = set(fm[fm < wstart].index)
    proj = sales[sales['角色'] == '项目'] if (scope == '项目' and '角色' in sales.columns) else sales
    rows = []
    for b in BRANDS:
        sub = proj[(proj['品牌'] == b) & (proj['患者ID'].isin(old))].copy()
        if sub.empty:
            continue
        visit = sub.groupby(['患者ID', 'ym'])['销售数量'].sum().reset_index()
        box_dist = visit['销售数量'].value_counts().sort_index()
        intervals = []
        for pid, g in sub.sort_values('销售时间').groupby('患者ID'):
            ts = g['销售时间'].sort_values().tolist()
            for i in range(1, len(ts)):
                intervals.append((ts[i] - ts[i - 1]).days)
        intervals = pd.Series(intervals)
        total = box_dist.sum()
        rows.append({
            '品种': b, '老患人数': sub['患者ID'].nunique(),
            '单次1盒占比': round(box_dist.get(1, 0) / total, 2) if total else '',
            '单次>=2盒占比': round(box_dist[box_dist.index >= 2].sum() / total, 2) if total else '',
            '单次>=3盒占比': round(box_dist[box_dist.index >= 3].sum() / total, 2) if total else '',
            '中位间隔天': round(intervals.median(), 0) if len(intervals) else '',
            '间隔>45天占比': round((intervals > 45).mean(), 2) if len(intervals) else ''
        })
    return pd.DataFrame(rows)

def repurchase_decomposition(sales, wstart=None, wend=None, wstart_prev=None, wend_prev=None):
    """复购率分解：全量 vs 老患，并估算新患稀释占比。"""
    if wstart is None or wend is None:
        wstart, wend, wstart_prev, wend_prev = resolve_window(sales, 'H1_2026')
    fm = sales.groupby('患者ID')['销售时间'].min()
    old = set(fm[fm < wstart].index)
    months = list(pd.period_range(wstart, wend, freq='M'))
    rows = []
    for b in BRANDS:
        sub = sales[sales['品牌'] == b]
        if sub.empty:
            continue
        r_all = _rolling_repurchase(sub, months, fm)
        pm = sub.groupby(['患者ID', 'ym']).size().reset_index()
        have = set(sub['ym'].unique())
        dil = []
        for m in months:
            if (m - 2) not in have or m not in have:
                continue
            bp = pm[pm['ym'] == m - 2]
            if bp.empty:
                continue
            fms = bp['患者ID'].map(fm).dt.to_period('M')
            dil.append((fms == m - 2).mean())
        dilution = np.mean(dil) if dil else np.nan
        r_old = _rolling_repurchase(sub[sub['患者ID'].isin(old)], months, fm)
        rows.append({
            '品种': b,
            '复购率_全_本期': round(r_all, 3) if pd.notna(r_all) else '',
            '新患稀释占比(估算)': round(dilution, 3) if pd.notna(dilution) else '',
            '老患占比': round(1 - dilution, 2) if pd.notna(dilution) else '',
            '复购率_老患_本期': round(r_old, 3) if pd.notna(r_old) else ''
        })
    return pd.DataFrame(rows)

def _split_doc(x):
    if pd.isna(x):
        return []
    return [p.strip() for p in str(x).replace('、', '|').replace('，', '|').replace(',', '|').split('|') if p.strip()]

def _newp_ratio(sub, fm, months):
    if sub.empty:
        return np.nan
    ids = [i for i in sub['患者ID'].unique() if pd.notna(i)]
    if not ids:
        return np.nan
    return fm.loc[ids].dt.to_period('M').isin(set(months)).mean()

def hospital_dimension(sales, scope, wstart=None, wend=None, wstart_prev=None, wend_prev=None):
    """医院维度：按医疗单位聚合 DOT、复购率、新患占比、主要来源药房。"""
    if wstart is None or wend is None:
        wstart, wend, wstart_prev, wend_prev = resolve_window(sales, 'H1_2026')
    months = list(pd.period_range(wstart, wend, freq='M'))
    dot_s, dot_e = _dot_window(wend)  # DOT 固定测量窗：窗口终点前推1年
    fm = sales.groupby('患者ID')['销售时间'].min()
    proj = sales[sales['角色'] == '项目'] if (scope == '项目' and '角色' in sales.columns) else sales
    rows = []
    for b in BRANDS:
        sub = proj[proj['品牌'] == b]
        if sub.empty:
            continue
        for hosp, g in sub.groupby('医疗单位'):
            n = g['患者ID'].nunique()
            if n < 5:
                continue
            d26 = _dot_snap(g, dot_s, dot_e)
            r26 = _rolling_repurchase(g, months, fm)
            ph_break = g.groupby('药房名称')['患者ID'].nunique().sort_values(ascending=False)
            contrib = ', '.join(f"{p}({nn})" for p, nn in ph_break.head(3).items())
            rows.append({
                '品种': b, '医疗单位': hosp, '患者数': n,
                'DOT_本期': round(d26, 2) if pd.notna(d26) else np.nan,
                '复购率_本期': round(r26, 3) if pd.notna(r26) else np.nan,
                '脱落率A': round(1 - r26, 3) if pd.notna(r26) else np.nan,
                '新患占比': round(_newp_ratio(g, fm, months), 3),
                '主要来源项目药房': ph_break.index[0] if len(ph_break) else '',
                '主要药房患者数': int(ph_break.iloc[0]) if len(ph_break) else '',
                '贡献项目药房明细': contrib
            })
    df = pd.DataFrame(rows)
    if df.empty:
        return df
    href = {b: df[df['品种'] == b]['DOT_本期'].median() for b in BRANDS}
    df['品种参考DOT'] = df['品种'].map(href)
    def flag(r):
        low = pd.notna(r['DOT_本期']) and r['DOT_本期'] < 0.5 * r['品种参考DOT']
        high_drop = pd.notna(r['脱落率A']) and r['脱落率A'] >= 0.4
        if high_drop and low:
            return '高脱落+低DOT-重点沟通'
        if high_drop:
            return '高脱落-重点沟通'
        if low and r['新患占比'] >= 0.5:
            return '新患致低DOT(结论)'
        if low:
            return '老患低DOT-真风险'
        if pd.notna(r['DOT_本期']) and r['DOT_本期'] < r['品种参考DOT']:
            return '低DOT-观察'
        return '正常'
    df['风险等级'] = df.apply(flag, axis=1)
    def action(r):
        if r['风险等级'] == '新患致低DOT(结论)':
            return '结论：低DOT由高新患占比稀释；动作：稳住新患+提升DOT；新患不足→反馈药企代表拓新'
        if '高脱落' in r['风险等级']:
            return '动作：院内外联动，针对该医院重点医生做观念沟通+续方管理'
        if '真风险' in r['风险等级']:
            return '动作：老患者依从性管理/续方提醒，核查外流'
        return '观察'
    df['结论与动作'] = df.apply(action, axis=1)
    return df.sort_values(['品种', '患者数'], ascending=[True, False])

def doctor_dimension(sales, scope, wstart=None, wend=None, wstart_prev=None, wend_prev=None):
    """医生维度：各项目药房下处方医生的患者量与 DOT。"""
    if wstart is None or wend is None:
        wstart, wend, wstart_prev, wend_prev = resolve_window(sales, 'H1_2026')
    dot_s, dot_e = _dot_window(wend)  # DOT 固定测量窗：窗口终点前推1年
    proj = sales[sales['角色'] == '项目'] if (scope == '项目' and '角色' in sales.columns) else sales
    proj2 = proj.copy()
    proj2['医生列表'] = proj2['处方医生'].map(_split_doc)
    proj2 = proj2.explode('医生列表')
    proj2 = proj2[proj2['医生列表'] != '']
    records = []
    for b in BRANDS:
        for ph in sorted(proj2[proj2['品牌'] == b]['药房名称'].unique()):
            sub = proj2[(proj2['品牌'] == b) & (proj2['药房名称'] == ph)]
            if sub.empty:
                continue
            st = sub.groupby('医生列表').agg(
                患者数=('患者ID', 'nunique'),
                主要医院=('医疗单位', lambda x: x.mode().iloc[0] if len(x.mode()) else '')
            )
            dmap = {doc: _dot_snap(g, dot_s, dot_e) for doc, g in sub.groupby('医生列表')}
            st['DOT_本期'] = pd.Series(dmap)
            st = st[st['患者数'] >= 1].sort_values('患者数', ascending=False)
            st['rank'] = range(1, len(st) + 1)
            for doc, row in st.iterrows():
                records.append({
                    '品种': b, '药房': ph, '城市': sales[sales['药房名称'] == ph]['城市'].dropna().iloc[0] if '城市' in sales.columns and not sales[sales['药房名称'] == ph]['城市'].dropna().empty else '',
                    'rank': int(row['rank']), '处方医生': doc,
                    '患者数': int(row['患者数']),
                    'DOT_本期': round(row['DOT_本期'], 2) if pd.notna(row['DOT_本期']) else np.nan,
                    '主要医院': row['主要医院']
                })
    doc_df = pd.DataFrame(records)
    if doc_df.empty:
        return doc_df, doc_df, doc_df
    dref = {b: doc_df[(doc_df['品种'] == b) & (doc_df['患者数'] >= 3)]['DOT_本期'].median() for b in BRANDS}
    doc_df['品种参考DOT'] = doc_df['品种'].map(dref)
    top1 = doc_df[doc_df['rank'] == 1].copy().rename(columns={
        '患者数': '药房该品种患者数', '处方医生': '最大患者量医生',
        'DOT_本期': '最大医生DOT'})
    top1 = top1[['品种', '药房', '城市', '药房该品种患者数', '最大患者量医生', '最大医生DOT', '主要医院']].sort_values(
        ['品种', '药房该品种患者数'], ascending=[True, False])
    top5 = doc_df[doc_df['rank'] <= 5].copy()
    top5['DOT相对参考'] = (top5['DOT_本期'] / top5['品种参考DOT']).round(2)
    top5 = top5[['品种', '药房', '城市', 'rank', '处方医生', '患者数', 'DOT_本期',
                 '品种参考DOT', 'DOT相对参考', '主要医院']].sort_values(['品种', '药房', 'rank'])
    docwatch = doc_df[(doc_df['rank'] <= 5) & (doc_df['患者数'] >= 5)].copy()
    docwatch = docwatch[docwatch['DOT_本期'] < docwatch['品种参考DOT']].copy()
    docwatch['DOT相对参考'] = (docwatch['DOT_本期'] / docwatch['品种参考DOT']).round(2)
    def newp_ratio_pharmacy(r):
        sub = sales[(sales['品牌'] == r['品种']) & (sales['药房名称'] == r['药房'])]
        return round(_newp_ratio(sub, sales.groupby('患者ID')['销售时间'].min(), list(pd.period_range(wstart, wend, freq='M'))), 3)
    docwatch['药房该品种新患占比'] = docwatch.apply(newp_ratio_pharmacy, axis=1)
    def dflag(r):
        if r['DOT相对参考'] < 0.5:
            return '明显偏低-重点沟通'
        if r['药房该品种新患占比'] >= 0.5:
            return '新患致低DOT(结论)'
        return '偏低-观察'
    docwatch['风险判定'] = docwatch.apply(dflag, axis=1)
    docwatch['结论与动作'] = np.where(
        docwatch['风险判定'] == '新患致低DOT(结论)',
        '结论：低DOT由新患多导致；动作：稳住新患+增DOT+反馈药企',
        np.where(docwatch['风险判定'] == '明显偏低-重点沟通',
                 '动作：重点医生观念沟通+患者依从性管理', '观察：结合右删失/样本判断'))
    docwatch = docwatch.sort_values(['品种', 'DOT_本期']).rename(columns={'DOT_本期': '该医生DOT'})
    docwatch = docwatch[['品种', '药房', '城市', 'rank', '处方医生', '患者数', '该医生DOT',
                         '品种参考DOT', 'DOT相对参考', '药房该品种新患占比', '风险判定', '结论与动作', '主要医院']]
    return doc_df, top1, top5, docwatch

def pharmacy_dimension(sales, wstart=None, wend=None, wstart_prev=None, wend_prev=None):
    """项目药房维度：大店低 DOT 识别。"""
    if wstart is None or wend is None:
        wstart, wend, wstart_prev, wend_prev = resolve_window(sales, 'H1_2026')
    months = list(pd.period_range(wstart, wend, freq='M'))
    dot_s, dot_e = _dot_window(wend)  # DOT 固定测量窗：窗口终点前推1年
    fm = sales.groupby('患者ID')['销售时间'].min()
    proj = sales[sales['角色'] == '项目'] if (scope == '项目' and '角色' in sales.columns) else sales
    rows = []
    for b in BRANDS:
        sub = proj[proj['品牌'] == b]
        if sub.empty:
            continue
        for ph, g in sub.groupby('药房名称'):
            n = g['患者ID'].nunique()
            if n < 5:
                continue
            d26 = _dot_snap(g, dot_s, dot_e)
            r26 = _rolling_repurchase(g, months, fm)
            rows.append({
                '品种': b, '药房': ph,
                '城市': sales[sales['药房名称'] == ph]['城市'].dropna().iloc[0] if '城市' in sales.columns and not sales[sales['药房名称'] == ph]['城市'].dropna().empty else '',
                '患者数': n,
                'DOT_本期': round(d26, 2) if pd.notna(d26) else np.nan,
                '复购率_本期': round(r26, 3) if pd.notna(r26) else np.nan,
                '脱落率A': round(1 - r26, 3) if pd.notna(r26) else np.nan,
                '新患占比': round(_newp_ratio(g, fm, months), 3)
            })
    df = pd.DataFrame(rows)
    if df.empty:
        return df
    pref = {b: df[df['品种'] == b]['DOT_本期'].median() for b in BRANDS}
    df['品种参考DOT'] = df['品种'].map(pref)
    vol_med = {b: df[df['品种'] == b]['患者数'].median() for b in BRANDS}
    df['体量大'] = df.apply(lambda r: r['患者数'] >= vol_med[r['品种']], axis=1)
    def pflag(r):
        if pd.isna(r['DOT_本期']):
            return '观察'
        if r['体量大'] and r['DOT_本期'] < 0.5 * r['品种参考DOT']:
            return '大店低DOT-重点改善'
        if r['体量大'] and r['DOT_本期'] < r['品种参考DOT']:
            return '大店低DOT-观察'
        if (not r['体量大']) and r['DOT_本期'] < 0.5 * r['品种参考DOT']:
            return '小店低DOT-观察'
        return '正常'
    df['风险等级'] = df.apply(pflag, axis=1)
    return df.sort_values(['品种', '患者数'], ascending=[True, False])

def drill_doctor_for_hospital(sales, scope, hosp_df, wstart=None, wend=None, wstart_prev=None, wend_prev=None):
    """钻取：异常医院下拖后腿的医生（按患者量 TOP5）。"""
    if wstart is None or wend is None:
        wstart, wend, wstart_prev, wend_prev = resolve_window(sales, 'H1_2026')
    proj = sales[sales['角色'] == '项目'] if (scope == '项目' and '角色' in sales.columns) else sales
    proj2 = proj.copy()
    proj2['医生列表'] = proj2['处方医生'].map(_split_doc)
    proj2 = proj2.explode('医生列表')
    proj2 = proj2[proj2['医生列表'] != '']
    rows = []
    if hosp_df.empty:
        return pd.DataFrame(rows)
    for _, hr in hosp_df[hosp_df['风险等级'].str.contains('重点沟通', na=False)].iterrows():
        b = hr['品种']
        hosp = hr['医疗单位']
        sub = proj2[(proj2['品牌'] == b) & (proj2['医疗单位'] == hosp)]
        if sub.empty:
            continue
        st = sub.groupby('医生列表').agg(
            患者数=('患者ID', 'nunique'),
            主要药房=('药房名称', lambda x: x.mode().iloc[0] if len(x.mode()) else '')
        )
        dmap = {doc: _dot_snap(g, dot_s, dot_e) for doc, g in sub.groupby('医生列表')}
        st['DOT_本期'] = pd.Series(dmap)
        st = st.sort_values('患者数', ascending=False).head(5)
        for doc, row in st.iterrows():
            rows.append({
                '品种': b, '异常医院': hosp, '医院风险等级': hr['风险等级'],
                '患者数': int(row['患者数']), '拖后腿医生': doc,
                '该医生DOT': round(row['DOT_本期'], 2) if pd.notna(row['DOT_本期']) else np.nan,
                '主要药房': row['主要药房'],
                '医院DOT': hr['DOT_本期'], '医院脱落率A': hr['脱落率A']
            })
    return pd.DataFrame(rows).sort_values(['品种', '异常医院', '患者数'], ascending=[True, True, False])

def communication_list(hosp_df, docwatch, pharm_df):
    """半年最该沟通清单：汇总医院/医生/药房层面的重点关注对象。"""
    comm = []
    if not hosp_df.empty:
        for _, r in hosp_df[hosp_df['风险等级'].isin(['高脱落-重点沟通', '高脱落+低DOT-重点沟通', '新患致低DOT(结论)'])].iterrows():
            comm.append({
                '维度': '医院', '品种': r['品种'], '对象': r['医疗单位'], '患者数': r['患者数'],
                '关键指标': f"DOT={r['DOT_本期']},脱落={r['脱落率A']}",
                '风险等级': r['风险等级'], '来源项目药房': r['主要来源项目药房'],
                '结论与动作': r['结论与动作']
            })
    if not docwatch.empty:
        for _, r in docwatch[docwatch['风险判定'].isin(['明显偏低-重点沟通', '新患致低DOT(结论)'])].iterrows():
            comm.append({
                '维度': '医生', '品种': r['品种'], '对象': f"{r['药房']}/{r['处方医生']}",
                '患者数': r['患者数'],
                '关键指标': f"DOT={r['该医生DOT']}(参考{r['品种参考DOT']})",
                '风险等级': r['风险判定'], '来源项目药房': r['药房'],
                '结论与动作': r['结论与动作']
            })
    if not pharm_df.empty:
        for _, r in pharm_df[pharm_df['风险等级'].str.startswith('大店低DOT', na=False)].iterrows():
            comm.append({
                '维度': '药房', '品种': r['品种'], '对象': r['药房'], '患者数': r['患者数'],
                '关键指标': f"DOT={r['DOT_本期']}(参考{r['品种参考DOT']})",
                '风险等级': r['风险等级'], '来源项目药房': r['药房'],
                '结论与动作': '大店/体量大但DOT偏低，优先改善：核查续方管理/外流/老患依从性'
            })
    return pd.DataFrame(comm).sort_values(['品种', '维度', '患者数'], ascending=[True, True, False])

def improvement_actions(res=None):
    """通用改进措施 + 结合本数据实际脱落原因的专项（动态）。归因维度仅用「细分类 + 一级分类」，不出现可控/不可控。
    res 含 crossref 时，按一级分类(医生相关/患者相关/其他)与细分原因追加专项。"""
    base = [
        {'方向': '稳住新患·提升DOT', '类型': '运营内部',
         '动作': '对高新患占比医院持续患者随访管理+患教，缩短新患首购→二购周期',
         '数据依据': '新患窗口 DOT 显著低于老患；成熟队列达二购率通常高于全量',
         '责任方': '患者服务运营+药房药师'},
        {'方向': '新患下降·反馈药企', '类型': '药企外部',
         '动作': '将各品种新患同比及"患者池仅靠老患"风险反馈药企区域代表，推动加强新患寻找',
         '数据依据': '新患月度趋势/项目药房新患变化',
         '责任方': '运营+药企代表'},
        {'方向': '高脱落医院·院内外联动', '类型': '院内外联动',
         '动作': '对脱落率>=0.4 的医院，联动对应重点医生做观念沟通+续方提醒',
         '数据依据': '医院维度高脱落+低DOT清单',
         '责任方': '药房+药企代表+医生'},
        {'方向': '老患低DOT·依从性管理', '类型': '运营内部',
         '动作': '对老患为主却 DOT 异常低的医生做依从性管理+多盒续方提醒，避免外流',
         '数据依据': '医生维度低 DOT 重点清单',
         '责任方': '药房药师'},
        {'方向': '大店低DOT·运营改善', '类型': '运营内部',
         '动作': '对体量大但 DOT 偏低的项目药房，核查续方管理、患者外流、随访执行',
         '数据依据': '项目药房风险等级',
         '责任方': '门店运营'},
    ]
    if res is None:
        return pd.DataFrame(base)
    # 动态专项1：品牌层面脱落原因的一级分类构成（取脱落率最高的几个品种）
    cb = res.get('crossref_brand')
    lvl_cols = {'医生': '一、医生相关原因数', '患者': '二、患者相关原因数', '其他': '三、其他原因数'}
    if cb is not None and not cb.empty:
        top = cb.dropna(subset=['销售脱落率B%']).sort_values('销售脱落率B%', ascending=False).head(3)
        for _, r in top.iterrows():
            br = r.get('品牌')
            if pd.isna(br):
                continue
            counts = {k: int(r.get(c, 0) or 0) for k, c in lvl_cols.items()}
            tot = sum(counts.values())
            if tot == 0:
                continue
            dom = max(counts, key=counts.get)
            dom_name = {'医生': '医生相关(医嘱停药/改间隔/换药/不良反应)',
                        '患者': '患者相关(经济/渠道分流/认知不足/随访缺失)',
                        '其他': '其他/未归类原因'}[dom]
            if dom == '医生':
                act = f'该品种脱落主因为「{dom_name}」，建专项：与区域/处方医生沟通方案连续性，明确停药原因并评估是否可逆'
                typ = '院内外联动'
            elif dom == '患者':
                act = f'该品种脱落主因为「{dom_name}」，建专项：续方提醒+支付/援助方案+患教+随访触达优化'
                typ = '运营内部'
            else:
                act = f'该品种脱落主因为「{dom_name}」，建专项：规范随访填写口径、提升原因可归类率'
                typ = '数据治理'
            base.append({
                '方向': f'{br} · 脱落原因干预',
                '类型': typ,
                '动作': act,
                '数据依据': f'销售脱落率B={r.get("销售脱落率B%")}%；随访脱落一级构成 医生{counts["医生"]}/患者{counts["患者"]}/其他{counts["其他"]}',
                '责任方': '患者服务运营+药房+医生' if typ == '院内外联动' else '患者服务运营+药房'})
    # 动态专项2：按药房×品种 脱落人数最多的细分原因 TOP5
    drp = res.get('dropout_reason_by_pharmacy')
    if drp is not None and not drp.empty:
        top = drp.sort_values('脱落患者数', ascending=False).head(5)
        for _, r in top.iterrows():
            n = int(r['脱落患者数'])
            if n <= 0:
                continue
            lvl = r.get('一级分类')
            base.append({
                '方向': f'{r["药房"]} · {r["品牌"]} · 脱落治理',
                '类型': lvl if lvl in ('一、医生相关原因', '二、患者相关原因', '三、其他原因') else '其他',
                '动作': f'该药房{r["品牌"]}有 {n} 名患者因「{r.get("脱落原因分类")}」脱落，按一级分类 {lvl} 落实针对性动作(见行动建议映射)',
                '数据依据': f'近似匹配随访原因={r.get("脱落原因分类")}，{n}人',
                '责任方': '对应药房药师' + ('+医生' if lvl == '一、医生相关原因' else '')})
    return pd.DataFrame(base)

def generate_word_report(res, sales_info=None):
    """生成可下载的 Word 复盘报告（关键结论 + 表格）。"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return None
    doc = Document()
    doc.add_heading('DTP 患者服务复盘报告', level=0)
    if sales_info:
        doc.add_paragraph(sales_info)
    # 窗口
    wi = res.get('window_info')
    if wi is not None and not wi.empty:
        r = wi.iloc[0]
        doc.add_paragraph(f"分析窗口：本期 {r['本期窗口起点']} ~ {r['本期窗口终点']}，"
                          f"同比 {r['上期窗口起点']} ~ {r['上期窗口终点']}")
    # 留存率
    doc.add_heading('一、留存率 Cohort', level=1)
    doc.add_paragraph('按首购月分群的新患队列，追踪后续每月仍在购药/在治的比例。')
    for title, key in [('口径1｜仅看购药时间', 'retention_overall'),
                       ('口径2｜结合说明书盒数覆盖', 'retention_cov_overall')]:
        df = res.get(key)
        if df is None or df.empty:
            continue
        doc.add_heading(title, level=2)
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Light Grid Accent 1'
        hdr = t.rows[0].cells
        for i, c in enumerate(df.columns):
            hdr[i].text = str(c)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if pd.isna(v) else str(v)
    # DOT
    doc.add_heading('二、DOT 分解（老患 vs 新患窗口）', level=1)
    dd = res.get('dot_decomposition')
    if dd is not None and not dd.empty:
        t = doc.add_table(rows=1, cols=len(dd.columns))
        t.style = 'Light Grid Accent 1'
        for i, c in enumerate(dd.columns):
            t.rows[0].cells[i].text = str(c)
        for _, row in dd.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if pd.isna(v) else str(v)
    # 脱落率B
    doc.add_heading('三、脱落率 B（累计沉默）', level=1)
    db = res.get('dropout_B_established_by_brand')
    if db is not None and not db.empty:
        t = doc.add_table(rows=1, cols=len(db.columns))
        t.style = 'Light Grid Accent 1'
        for i, c in enumerate(db.columns):
            t.rows[0].cells[i].text = str(c)
        for _, row in db.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if pd.isna(v) else str(v)
    # 脱落原因分类（修正后·细分类框架）
    drd = res.get('dropout_reason_detail')
    if drd is not None and not drd.empty:
        doc.add_heading('三之二、脱落原因分类（修正后·细分类框架）', level=1)
        doc.add_paragraph('按 churn_logic 框架（医生相关→患者相关→其他，统一兜底）对随访原始原因清洗归类；'
                          '下方为「有记载脱落原因」的记录分布。')
        t = doc.add_table(rows=1, cols=len(drd.columns))
        t.style = 'Light Grid Accent 1'
        for i, c in enumerate(drd.columns):
            t.rows[0].cells[i].text = str(c)
        for _, row in drd.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if pd.isna(v) else str(v)
        dmeta = res.get('dropout_reason_meta')
        if isinstance(dmeta, dict):
            doc.add_paragraph(
                f"覆盖提示：随访总记录 {dmeta.get('随访总记录')}，其中「有记载原因」{dmeta.get('有原因记录')} 条，"
                f"「空原因(≠无脱落，可能是随访员未填)」{dmeta.get('无原因记录')} 条"
                f"（占 {dmeta.get('无原因占比%')}%）。空原因已单独计数，未混入任何细类。")
    # 沟通清单
    doc.add_heading('四、重点关注清单', level=1)
    comm = res.get('communication_list')
    if comm is not None and not comm.empty:
        t = doc.add_table(rows=1, cols=len(comm.columns))
        t.style = 'Light Grid Accent 1'
        for i, c in enumerate(comm.columns):
            t.rows[0].cells[i].text = str(c)
        for _, row in comm.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if pd.isna(v) else str(v)
    else:
        doc.add_paragraph('当前数据未触发重点关注阈值。')
    # 改进措施
    doc.add_heading('五、改进措施与责任', level=1)
    ia = res.get('improvement_actions')
    if ia is not None and not ia.empty:
        t = doc.add_table(rows=1, cols=len(ia.columns))
        t.style = 'Light Grid Accent 1'
        for i, c in enumerate(ia.columns):
            t.rows[0].cells[i].text = str(c)
        for _, row in ia.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if pd.isna(v) else str(v)
    # 口径说明
    doc.add_heading('六、口径说明', level=1)
    doc.add_paragraph(
        'DOT = 窗口内购买总盒数 / 去重患者数（盒/人），不换算月数；老患=窗口起始前首购，新患窗口=首购落在窗口内。'
        '留存率口径1仅看当月是否有购药，口径2按说明书每盒天数×购买盒数计算覆盖期。'
        '脱落率B = 末次购药距数据终点 > N×用药间隔。')
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ============================ 编排 ============================
def run_analysis(sales, followup=None, max_k=12, mult=3, with_patient_crossref=True, preset='H1_2026', custom_start=None, custom_end=None, scope='项目'):
    result = {}
    # 分析范围（项目药房 / 全部药房）：TOP清单判定的 角色 列，按 scope 过滤底表
    if scope == '项目' and '角色' in sales.columns:
        sales_scope = sales[sales['角色'] == '项目'].copy()
    else:
        sales_scope = sales.copy()
    # 窗口信息（统一时间窗）
    wstart, wend, wstart_prev, wend_prev = resolve_window(sales, preset, custom_start, custom_end)
    if preset == 'custom':
        _preset_label = f'自定义 {wstart.date()}~{wend.date()}'
    else:
        _preset_label = preset
    dot_s, dot_e = _dot_window(wend)
    result['window_info'] = pd.DataFrame([{
        '时间窗预设': _preset_label,
        '本期窗口起点': wstart.date() if wstart is not None else None,
        '本期窗口终点': wend.date() if wend is not None else None,
        '上期窗口起点': wstart_prev.date() if wstart_prev is not None else None,
        '上期窗口终点': wend_prev.date() if wend_prev is not None else None,
        'DOT口径(盒/人,不按月归一)': f'窗口终点前推1年固定1年窗：{dot_s.date()}~{dot_e.date()}' if dot_s is not None else '全量(无终点)',
        '分析范围': scope,
    }])
    # 留存率等月度指标：按时间窗过滤销售（full 预设 = 全量），并用 scope 限定的底表
    sales_win = sales_scope.copy()
    if wstart is not None and wend is not None:
        sales_win = sales[(sales['销售时间'] >= wstart) & (sales['销售时间'] <= wend)].copy()
    # 留存率 口径1：仅看购药时间（当月是否有购药）
    result['retention_overall'] = compute_retention(sales_win, max_k)
    result['retention_by_brand'] = compute_retention_by_brand(sales_win, max_k)
    # 分品种×药房留存率：保证 key 始终存在（个别数据异常时降级为空表，不丢 key 不崩溃）
    try:
        result['retention_by_brand_pharmacy'] = compute_retention_by_brand_pharmacy(sales_win, max_k)
    except Exception as _e:
        result['retention_by_brand_pharmacy'] = pd.DataFrame()
    # 留存率 口径2：结合说明书盒数覆盖期
    result['retention_cov_overall'] = compute_retention_coverage(sales_win, max_k)
    result['retention_cov_by_brand'] = compute_retention_coverage_by_brand(sales_win, max_k)
    try:
        result['dropout_A'] = compute_dropout_A(sales_win)
    except Exception:
        result['dropout_A'] = {}
    last = pd.DataFrame()  # 默认空，确保下游 groupby 不报错
    try:
        last = compute_dropout_B(sales_scope, mult, end=wend)
        result['dropout_B_patient'] = last
    except Exception:
        result['dropout_B_patient'] = last
    # 脱落率B 分品种（全部患者）
    if not last.empty and '品牌' in last.columns and '脱落B' in last.columns:
        bbrand = last.groupby('品牌')['脱落B'].agg(脱落患者数='sum', 患者总数='count')
        bbrand['脱落率B%'] = (bbrand['脱落患者数'] / bbrand['患者总数'] * 100).round(1)
        result['dropout_B_by_brand'] = bbrand.reset_index().sort_values('脱落率B%', ascending=False)
    else:
        result['dropout_B_by_brand'] = pd.DataFrame()
    # 脱落率B 分品种（已观察患者，剔除新近右删失）
    if not last.empty and 'established' in last.columns and '脱落B' in last.columns:
        est = last[last['established']]
        ebrand = est.groupby('品牌')['脱落B'].agg(脱落患者数_已观察='sum', 患者总数_已观察='count')
        ebrand['脱落率B_已观察%'] = (ebrand['脱落患者数_已观察'] / ebrand['患者总数_已观察'] * 100).round(1)
        result['dropout_B_established_by_brand'] = ebrand.reset_index().sort_values('脱落率B_已观察%', ascending=False)
    else:
        result['dropout_B_established_by_brand'] = pd.DataFrame()
    # 复购率B 分品种（结合用药周期 / 首购→二购 on-cycle），镜像脱落率B 结构
    rb_pat = pd.DataFrame()
    try:
        rb_pat = compute_repurchase_B(sales_scope, mult, end=wend)
        result['repurchase_B_patient'] = rb_pat
    except Exception:
        result['repurchase_B_patient'] = rb_pat
    try:
        if not rb_pat.empty:
            rbb = repurchase_B_by_brand(sales_scope, mult, end=wend)
            result['repurchase_B_by_brand'] = rbb
            if not rbb.empty:
                result['repurchase_B_established_by_brand'] = rbb[['品种', '复购率B_已观察%', '已观察_按周期复购数', '已观察患者数']].copy()
            else:
                result['repurchase_B_established_by_brand'] = pd.DataFrame()
        else:
            result['repurchase_B_by_brand'] = pd.DataFrame()
            result['repurchase_B_established_by_brand'] = pd.DataFrame()
    except Exception:
        # 任何内部异常都兜底写入空表，保证下游 app.py 不会 KeyError 崩溃
        result['repurchase_B_by_brand'] = pd.DataFrame()
        result['repurchase_B_established_by_brand'] = pd.DataFrame()
    result['action_map'] = pd.DataFrame(ACTION_MAP, columns=['一级分类', '典型原因', '建议动作', '责任方'])
    # 用于脱落原因统计的窗内随访（避免重复过滤）
    if followup is not None and not followup.empty:
        fw = _reason_windowed(followup, wstart, wend)
        try:
            result['crossref_brand'] = crossref_brand(last, followup, wstart, wend)
        except Exception:
            result.pop('crossref_brand', None)
        # 脱落原因细分类分布（总表 + 一级汇总 + 按品种拆分 + 覆盖）
        try:
            d_dist, d_lvl1, d_meta = dropout_reason_distribution(followup, wstart, wend)
            result['dropout_reason_detail'] = d_dist
            result['dropout_reason_lvl1'] = d_lvl1
            result['dropout_reason_meta'] = d_meta
            result['dropout_reason_by_brand'] = dropout_reason_detail_by_brand(followup, wstart, wend)
        except Exception:
            result.pop('dropout_reason_detail', None)
        # 患者级双视角明细（随访侧：最新一次 vs 最近一次有原因），窗内
        try:
            _fu_meta = {
                'patient_col': '患者姓名', 'store_col': '药房原始',
                'reason_col': '原因原始', 'status_col': '用药周期状态',
                'date_cols': [c for c in cl.DATE_COL_CANDIDATES if c in followup.columns],
            }
            result['patient_level_detail'] = cl.patient_level_detail(fw, _fu_meta)
        except Exception:
            result['patient_level_detail'] = pd.DataFrame()
        # 自动发现药房名对齐候选（写 pharm_map.auto.json，供人工确认后抄入 pharm_map.json）
        try:
            result['pharm_map_suggestion'] = suggest_pharm_map(sales, followup)
        except Exception:
            result['pharm_map_suggestion'] = {}
        if with_patient_crossref:
            pm = load_pharm_map()
            try:
                cp = crossref_patient(last, followup, pm, wstart, wend)
                result['crossref_patient'] = cp
            except Exception:
                result.pop('crossref_patient', None)
                cp = None
            # 覆盖率概览（含低匹配报警）
            if cp is not None and not cp.empty:
                try:
                    result['crossref_coverage'] = crossref_coverage_summary(cp)
                except Exception:
                    result.pop('crossref_coverage', None)
                # 脱落原因 × 药房 × 品种 × 细类（仅已匹配且有原因记录）
                matched = cp[(cp['关联方式'] != '无随访匹配') & cp['脱落原因分类'].notna()]
                if not matched.empty:
                    rc = matched.groupby(['药房', '品牌', '脱落原因分类', '一级分类']).size().reset_index(name='脱落患者数')
                    result['dropout_reason_by_pharmacy'] = rc
    # ===== 新增：DOT / 新患 / 复购 / 维度分析（逐个防御，单点失败不拖垮整体） =====
    def _safe_df(key, fn, *a):
        try:
            result[key] = fn(*a)
        except Exception:
            result[key] = pd.DataFrame()
    _safe_df('dot_decomposition', dot_decomposition, sales_scope, wstart, wend, wstart_prev, wend_prev)
    _safe_df('new_patient_monthly', new_patient_monthly, sales_scope)
    _safe_df('new_patient_pharmacy_decline', new_patient_pharmacy_decline, sales, scope, wstart, wend, wstart_prev, wend_prev)
    _safe_df('old_patient_multi_box', old_patient_multi_box, sales, scope, wstart, wend, wstart_prev, wend_prev)
    _safe_df('repurchase_decomposition', repurchase_decomposition, sales_scope, wstart, wend, wstart_prev, wend_prev)
    _safe_df('hospital_dimension', hospital_dimension, sales, scope, wstart, wend, wstart_prev, wend_prev)
    try:
        doc_df, top1, top5, docwatch = doctor_dimension(sales, scope, wstart, wend, wstart_prev, wend_prev)
        result['doctor_all'] = doc_df
        result['doctor_top1'] = top1
        result['doctor_top5'] = top5
        result['doctor_low_dot_watch'] = docwatch
    except Exception:
        result['doctor_all'] = pd.DataFrame()
        result['doctor_top1'] = pd.DataFrame()
        result['doctor_top5'] = pd.DataFrame()
        result['doctor_low_dot_watch'] = pd.DataFrame()
    _safe_df('pharmacy_dimension', pharmacy_dimension, sales, scope, wstart, wend, wstart_prev, wend_prev)
    try:
        result['drill_hospital_doctor'] = drill_doctor_for_hospital(sales, scope, result['hospital_dimension'], wstart, wend, wstart_prev, wend_prev)
    except Exception:
        result['drill_hospital_doctor'] = pd.DataFrame()
    try:
        result['communication_list'] = communication_list(
            result['hospital_dimension'], result['doctor_low_dot_watch'], result['pharmacy_dimension'])
    except Exception:
        result['communication_list'] = pd.DataFrame()
    try:
        result['improvement_actions'] = improvement_actions(result)
    except Exception:
        result['improvement_actions'] = pd.DataFrame()
    return result

if __name__ == '__main__':
    import sys
    SRC = r'F:/厂家项目/阿斯利康/2026/半年复盘/英飞凡、利普卓、泰瑞沙、优赫得6.26.xlsx'
    F25 = 'E:/下载内容/历史任务 - 2026-07-22T105636.443.xls'
    F26 = 'E:/下载内容/历史任务 - 2026-07-22T105440.252.xls'
    sales = load_sales(SRC)
    followup = load_followup([F25, F26])
    res = run_analysis(sales, followup)
    print('留存率-口径1(仅购药时间) 整体:')
    print(res['retention_overall'].head(8).to_string(index=False))
    print('\n留存率-口径2(结合盒数覆盖) 整体:')
    print(res['retention_cov_overall'].head(8).to_string(index=False))
    print('\n自丰富映射表 aliases 条数:', len(_ALIAS['aliases']), ' unresolved:', _ALIAS['unresolved'])
    print('\n脱落率A 整体月度:')
    print(res['dropout_A']['整体_月度'].tail(8).to_string(index=False))
    print('\n脱落率B 分品种:')
    print(res['dropout_B_by_brand'].to_string(index=False))
    if 'crossref_brand' in res:
        print('\n跨表关联(品牌层面):')
        print(res['crossref_brand'].to_string(index=False))
    if 'crossref_patient' in res:
        cp = res['crossref_patient']
        print('\n近似患者级关联匹配率: {}/{} = {:.1f}%'.format(
            cp.attrs.get('matched', 0), cp.attrs.get('total', 0),
            cp.attrs.get('matched', 0) / max(cp.attrs.get('total', 1), 1) * 100))
